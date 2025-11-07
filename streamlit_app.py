import streamlit as st 
import pandas as pd
from io import BytesIO
from docx import Document
from datetime import date as dt_date
import gspread
from google.oauth2.service_account import Credentials
import altair as alt
import re
import matplotlib.pyplot as plt

# Hide Streamlit UI elements
hide_streamlit_ui = """
    <style>
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    header {visibility: hidden;}
    .stDeployButton {display: none;}
    div[data-testid="stDecoration"] {visibility: hidden;}
    [data-testid="stToolbar"] {visibility: hidden !important;}
    [data-testid="stStatusWidget"] {visibility: hidden !important; height: 0; overflow: hidden;}
    </style>
"""
st.markdown(hide_streamlit_ui, unsafe_allow_html=True)


# Custom CSS for attractive design
custom_css = """
    <style>
    :root {
        --primary-color: #1f77b4;
        --secondary-color: #ff7f0e;
        --success-color: #2ca02c;
        --danger-color: #d62728;
        --bg-light: #f8f9fa;
        --border-color: #e0e0e0;
    }

    * {
        margin: 0;
        padding: 0;
    }

    .main-header {
        background: linear-gradient(135deg, #1f77b4 0%, #0d5a9a 100%);
        color: white;
        padding: 2.5rem 2rem;
        border-radius: 12px;
        margin-bottom: 2rem;
        box-shadow: 0 4px 15px rgba(31, 119, 180, 0.3);
    }

    .main-header h1 {
        font-size: 2.5rem;
        font-weight: 700;
        margin-bottom: 0.5rem;
        text-shadow: 0 2px 4px rgba(0, 0, 0, 0.1);
    }

    .main-header p {
        font-size: 1rem;
        opacity: 0.95;
        font-weight: 300;
    }

    .form-section {
        background: white;
        border-radius: 12px;
        padding: 2rem;
        margin-bottom: 2rem;
        border: 1px solid var(--border-color);
        box-shadow: 0 2px 8px rgba(0, 0, 0, 0.06);
    }

    .form-section h2 {
        color: #1f77b4;
        margin-bottom: 1.5rem;
        font-size: 1.5rem;
        border-bottom: 3px solid #1f77b4;
        padding-bottom: 0.5rem;
    }

    .data-section {
        background: white;
        border-radius: 12px;
        padding: 2rem;
        margin-bottom: 2rem;
        border: 1px solid var(--border-color);
        box-shadow: 0 2px 8px rgba(0, 0, 0, 0.06);
    }

    .data-section h2 {
        color: #1f77b4;
        margin-bottom: 1.5rem;
        font-size: 1.5rem;
        border-bottom: 3px solid #1f77b4;
        padding-bottom: 0.5rem;
    }

    .table-container {
        overflow-x: auto;
        border-radius: 8px;
        border: 1px solid var(--border-color);
        margin-bottom: 1.5rem;
        max-height: 500px;
        overflow-y: auto;
    }

    .table-container table {
        width: 100%;
        border-collapse: collapse;
    }

    .table-container thead th {
        background: linear-gradient(135deg, #1f77b4 0%, #0d5a9a 100%);
        color: white;
        padding: 1rem;
        text-align: left;
        font-weight: 600;
        position: sticky;
        top: 0;
        z-index: 10;
    }

    .table-container tbody td {
        padding: 0.75rem 1rem;
        border-bottom: 1px solid var(--border-color);
    }

    .table-container tbody tr:hover {
        background-color: #f0f7ff;
        transition: background-color 0.2s ease;
    }

    .table-container tbody tr:nth-child(even) {
        background-color: #fafbfc;
    }

    .download-section {
        display: flex;
        gap: 1rem;
        margin-top: 1.5rem;
        flex-wrap: wrap;
    }

    .stButton > button {
        background: linear-gradient(135deg, #1f77b4 0%, #0d5a9a 100%) !important;
        color: white !important;
        border: none !important;
        border-radius: 8px !important;
        padding: 0.75rem 1.5rem !important;
        font-weight: 600 !important;
        transition: all 0.3s ease !important;
        box-shadow: 0 4px 10px rgba(31, 119, 180, 0.2) !important;
    }

    .stButton > button:hover {
        box-shadow: 0 6px 15px rgba(31, 119, 180, 0.4) !important;
        transform: translateY(-2px) !important;
    }

    .stForm {
        border: none !important;
    }

    .stSelectbox, .stTextInput, .stDateInput {
        margin-bottom: 1rem;
    }

    .stAlert {
        border-radius: 8px !important;
        margin-bottom: 1rem;
    }

    .diameter-label {
        font-weight: 600;
        color: #333;
        margin-top: 0.5rem;
    }

    .info-box {
        background: #e3f2fd;
        border-left: 4px solid #1f77b4;
        padding: 1rem;
        border-radius: 6px;
        margin-bottom: 1rem;
    }

    .download-section {
        display: flex;
        gap: 1rem;
        flex-wrap: wrap;
        padding-top: 1rem;
        border-top: 1px solid var(--border-color);
    }

    .page-controls {
        display: flex;
        align-items: center;
        gap: 1rem;
        margin-bottom: 1rem;
    }

    body {
        background-color: #f5f7fa;
    }
    </style>
"""
st.markdown(custom_css, unsafe_allow_html=True)

st.set_page_config(layout="wide", page_title="Roll Profile Data Entry")

# --- Google Sheets Config ---
SHEET_NAME = "Roll_Data"
SCOPE = [
    "https://www.googleapis.com/auth/spreadsheets",
    "https://www.googleapis.com/auth/drive"
]

creds_dict = st.secrets["gcp_service_account"]
creds = Credentials.from_service_account_info(creds_dict, scopes=SCOPE)
client = gspread.authorize(creds)
sheet = client.open(SHEET_NAME).sheet1

# --- Roll Config ---
DISTANCES = [100, 350, 600, 850, 1100, 1350, 1600]
MIN_DIA = 1245.0
MAX_DIA = 1352.0

# --- Header ---
st.markdown("""
    <div class="main-header">
        <h1>📊 Backup Roll Profile Data Entry</h1>
        <p>Manage and track roll specifications with ease</p>
    </div>
""", unsafe_allow_html=True)

# Load existing data
existing_data = sheet.get_all_records()
df = pd.DataFrame(existing_data)

# --- Entry Form ---
form_diameters = {}
with st.container():
    st.markdown('<div class="form-section">', unsafe_allow_html=True)
    with st.form("entry_form", clear_on_submit=False):
        st.markdown("### ➕ Add New Roll Entry")
        
        col1, col2, col3 = st.columns(3)
        with col1:
            entry_date = st.date_input("📅 Date", value=dt_date.today())
        with col2:
            roll_no = st.text_input("🏷️ Roll No (required)").strip().upper()
        with col3:
            stand = st.selectbox(" Stand", ['Select', 'F1', 'F2', 'F3', 'F4', 'F5', 'F6', 'ROUGHING', 'DC'], index=0)

        col1, col2 = st.columns(2)
        with col1:
            position = st.selectbox("📍 Position", ['Select', 'TOP', 'BOTTOM'], index=0)
        with col2:
            crown = st.selectbox(" Crown", ['Select', 'STRAIGHT', '+100µ', '+200µ'], index=0)

        st.markdown('<p class="diameter-label">📏 Diameters (mm) — must be between 1245 and 1352</p>', unsafe_allow_html=True)
        
        # Single column for diameter inputs
        for d in DISTANCES:
            val = st.text_input(f"{d} mm", value="", key=f"dia_{d}", placeholder="Enter value")
            try:
                form_diameters[d] = float(val) if val.strip() != "" else 0
            except ValueError:
                form_diameters[d] = 0

        submitted = st.form_submit_button("💾 Save Entry", use_container_width=True)

    st.markdown('</div>', unsafe_allow_html=True)

# --- Save Entry ---
if submitted:
    errors = []

    if roll_no == "":
        errors.append("❌ Roll No cannot be empty")
    
    if stand == "Select":
        errors.append("❌ Please select a Stand")
    
    if position == "Select":
        errors.append("❌ Please select a Position")
    
    if crown == "Select":
        errors.append("❌ Please select a Crown type")

    filtered_diameters = {}
    for d, v in form_diameters.items():
        if v == 0:
            continue
        if not (MIN_DIA <= v <= MAX_DIA):
            errors.append(f"❌ {d} mm value {v} out of range [{MIN_DIA}-{MAX_DIA}]")
        else:
            filtered_diameters[d] = v

    if errors:
        for e in errors:
            st.error(e)
    else:
        row = [str(entry_date), roll_no, stand, position, crown] + [filtered_diameters.get(d, "") for d in DISTANCES]
        sheet.append_row(row)
        st.success(f"✅ Entry saved for Roll No: {roll_no}")

        existing_data = sheet.get_all_records()
        df = pd.DataFrame(existing_data)

# --- Show Data ---
with st.container():
    st.markdown('<div class="data-section">', unsafe_allow_html=True)
    st.markdown("### 📋 Stored Data")
    
    if df.empty:
        st.markdown('<div class="info-box">📭 No entries yet. Start by adding a new roll entry above.</div>', unsafe_allow_html=True)
    else:
        for col in df.columns:
            if df[col].dtype in ["float64", "int64"]:
                df[col] = df[col].map(lambda x: f"{x:.2f}" if x != "" else "")

        df_display = df.reset_index(drop=True)

        # Pagination
        page_size = 10
        total_pages = (len(df_display) - 1) // page_size + 1
        
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            page = st.number_input("📄 Page", min_value=1, max_value=total_pages, step=1, label_visibility="collapsed")
        
        start = (page - 1) * page_size
        end = start + page_size

        # Display table with custom scrolling
        st.markdown('<div class="table-container">', unsafe_allow_html=True)
        st.dataframe(df_display.iloc[start:end], use_container_width=True, hide_index=True)
        st.markdown('</div>', unsafe_allow_html=True)

        st.markdown(f"<p style='text-align: center; color: #666; font-size: 0.9rem; margin: 1rem 0;'>Page {page} of {total_pages} | Total entries: {len(df_display)}</p>", unsafe_allow_html=True)

    # --- Download Functions ---
    def to_excel_bytes(df):
        output = BytesIO()
        df.to_excel(output, index=False, sheet_name="RollData")
        output.seek(0)
        return output.getvalue()

    def to_word_bytes(df):
        doc = Document()
        doc.add_heading("Roll Profile Data", level=1)
        table = doc.add_table(rows=1, cols=len(df.columns))
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for i, col in enumerate(df.columns):
            hdr[i].text = str(col)
        for _, r in df.iterrows():
            cells = table.add_row().cells
            for j, col in enumerate(df.columns):
                cells[j].text = str(r[col])
        out = BytesIO()
        doc.save(out)
        out.seek(0)
        return out.getvalue()

    # --- Download Buttons ---
    if not df.empty:
        st.markdown('<div class="download-section">', unsafe_allow_html=True)
        col1, col2 = st.columns(2)
        with col1:
            st.download_button(
                "⬇️ Download Excel",
                data=to_excel_bytes(df),
                file_name="roll_data.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True
            )
        with col2:
            st.download_button(
                "⬇️ Download Word",
                data=to_word_bytes(df),
                file_name="roll_data.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
        st.markdown('</div>', unsafe_allow_html=True)

    st.markdown('</div>', unsafe_allow_html=True)

# ---------- Plot Roll Profile Section ----------
st.markdown('<div class="data-section">', unsafe_allow_html=True)
st.markdown("## 📈 Plot Roll Profile")

if df.empty:
    st.info("No data to plot.")
else:
    # --- Helper functions ---
    def find_col_by_candidates(col_list, candidates):
        cols_map = {c.strip().lower(): c for c in col_list}
        for cand in candidates:
            if cand.strip().lower() in cols_map:
                return cols_map[cand.strip().lower()]
        return None

    # Clean column headers
    df_plot = df.copy()
    df_plot.rename(columns={c: c.strip() for c in df_plot.columns}, inplace=True)
    norm_cols = list(df_plot.columns)

    # Find key columns
    date_col = find_col_by_candidates(norm_cols, ["date", "entry date", "entry_date"])
    roll_col = find_col_by_candidates(norm_cols, ["roll no", "rollno", "roll_no", "roll"])

    if date_col is None or roll_col is None:
        st.error("Could not find required 'Date' or 'Roll No' columns in sheet.")
    else:
        # Format date labels
        try:
            df_plot[date_col] = pd.to_datetime(df_plot[date_col])
            df_plot["_date_label"] = df_plot[date_col].dt.strftime("%Y-%m-%d")
        except Exception:
            df_plot["_date_label"] = df_plot[date_col].astype(str)

        # Detect distance columns
        desired_distances = [100, 350, 600, 850, 1100, 1350, 1600]
        found_distance_cols = []
        for col in norm_cols:
            m = re.search(r"(\d+)", str(col))
            if m:
                try:
                    dist = int(m.group(1))
                except Exception:
                    dist = None
                if dist in desired_distances:
                    found_distance_cols.append((dist, col))

        found_distance_cols = sorted(
            found_distance_cols, key=lambda x: desired_distances.index(x[0])
        )

        if not found_distance_cols:
            st.error("No distance columns (100,350,...) found in sheet.")
        else:
            # Roll selection
            roll_options = sorted(df_plot[roll_col].astype(str).unique())
            selected_roll = st.selectbox("Select Roll No", ["-- choose --"] + roll_options)

            if selected_roll and selected_roll != "-- choose --":
                roll_rows = df_plot[df_plot[roll_col].astype(str) == str(selected_roll)].copy()
                if roll_rows.empty:
                    st.warning("No rows for that Roll No.")
                else:
                    date_options = roll_rows["_date_label"].tolist()
                    default_dates = [date_options[-1]] if date_options else []
                    chosen_dates = st.multiselect(
                        "Select one or more Dates to plot (multiple lines)",
                        options=date_options,
                        default=default_dates,
                    )

                    if not chosen_dates:
                        st.info("Select at least one date to plot.")
                    else:
                        # Build long-form dataframe
                        rows = []
                        for _, r in roll_rows.iterrows():
                            label = r["_date_label"]
                            if label not in chosen_dates:
                                continue
                            for d, colname in found_distance_cols:
                                raw = r.get(colname, None)
                                try:
                                    if raw is None or str(raw).strip() == "":
                                        val = None
                                    else:
                                        val = float(str(raw).strip().replace(",", ""))
                                except Exception:
                                    val = None
                                if val is not None:
                                    rows.append(
                                        {"DateLabel": label, "Distance": int(d), "Diameter": val}
                                    )

                        if not rows:
                            st.warning("No numeric data available for selected dates.")
                        else:
                            plot_df = pd.DataFrame(rows).sort_values(["DateLabel", "Distance"])

                            # Chart settings
                            min_dist = int(plot_df["Distance"].min())
                            max_dist = int(plot_df["Distance"].max())
                            y_min = float(plot_df["Diameter"].min())
                            y_max = float(plot_df["Diameter"].max())
                            y_pad = (y_max - y_min) * 1 if (y_max - y_min) > 0 else 0.6
                            y_domain = [y_min - y_pad, y_max + y_pad]
                            x_axis_values = [d for d, _ in found_distance_cols]

                            # Altair chart
                            chart = (
                                alt.Chart(plot_df, title="Dirty Roll Profile")
                                .mark_line(
                                    point=alt.OverlayMarkDef(filled=True, size=60),
                                    interpolate="monotone",
                                )
                                .encode(
                                    x=alt.X(
                                        "Distance:Q",
                                        title="Distance (mm)",
                                        scale=alt.Scale(domain=[min_dist, max_dist]),
                                        axis=alt.Axis(values=x_axis_values),
                                    ),
                                    y=alt.Y(
                                        "Diameter:Q",
                                        title="Diameter (mm)",
                                        scale=alt.Scale(domain=y_domain),
                                    ),
                                    color=alt.Color("DateLabel:N", title="Date"),
                                    tooltip=[
                                        alt.Tooltip("DateLabel", title="Date"),
                                        alt.Tooltip("Distance", title="Distance (mm)"),
                                        alt.Tooltip("Diameter", title="Diameter (mm)", format=".3f"),
                                    ],
                                )
                                .properties(height=380)
                            )

                            st.altair_chart(chart, use_container_width=True)

                            # Display data table below chart
                            st.markdown("**Plotted Roll Data :**")
                            display_df = plot_df[["Distance", "Diameter"]].copy()
                            display_df = display_df.sort_values("Distance").reset_index(drop=True)
                            st.dataframe(display_df, use_container_width=True, hide_index=True)

                            # Download chart as Excel with embedded chart
                            def to_chart_excel_bytes(plot_data, roll_id, dates_selected):
                                output = BytesIO()
                                
                                try:
                                    # Try xlsxwriter first for chart support
                                    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
                                        workbook = writer.book
                                        worksheet = workbook.add_worksheet('Roll Profile')
                                        writer.sheets['Roll Profile'] = worksheet
                                        
                                        # Formats
                                        title_format = workbook.add_format({
                                            'bold': True, 
                                            'font_size': 16, 
                                            'align': 'center',
                                            'valign': 'vcenter'
                                        })
                                        info_format = workbook.add_format({'bold': True, 'font_size': 11})
                                        data_format = workbook.add_format({'font_size': 10})
                                        header_format = workbook.add_format({
                                            'bold': True, 
                                            'bg_color': '#1f77b4', 
                                            'font_color': 'white', 
                                            'align': 'center'
                                        })
                                        
                                        # Title - using Roll ID
                                        worksheet.merge_range('A1:C1','Dirty roll profile', title_format)
                                        worksheet.set_row(0, 25)
                                        
                                        # Roll information
                                        row = 2
                                        worksheet.write(row, 0, 'Roll No:', info_format)
                                        worksheet.write(row, 1, roll_id, data_format)
                                        
                                        row += 1
                                        worksheet.write(row, 0, 'Date(s):', info_format)
                                        worksheet.write(row, 1, ', '.join(dates_selected), data_format)
                                        
                                        # Prepare data organized by date
                                        row += 2
                                        start_row = row
                                        
                                        # Write headers
                                        worksheet.write(row, 0, 'Distance', header_format)
                                        
                                        dates_list = sorted(plot_data['DateLabel'].unique())
                                        for idx, date_label in enumerate(dates_list):
                                            worksheet.write(row, idx + 1, 'Diameter', header_format)
                                        
                                        row += 1
                                        data_start_row = row
                                        
                                        # Write distance and diameter data
                                        distances = sorted(plot_data['Distance'].unique())
                                        for dist in distances:
                                            worksheet.write(row, 0, dist, data_format)
                                            for idx, date_label in enumerate(dates_list):
                                                dia_val = plot_data[
                                                    (plot_data['Distance'] == dist) & 
                                                    (plot_data['DateLabel'] == date_label)
                                                ]['Diameter'].values
                                                if len(dia_val) > 0:
                                                    worksheet.write(row, idx + 1, dia_val[0], data_format)
                                            row += 1
                                        
                                        data_end_row = row - 1
                                        
                                        # Create chart
                                        chart = workbook.add_chart({'type': 'line'})
                                        
                                        # Add series for each date
                                        for idx, date_label in enumerate(dates_list):
                                            col_letter = chr(66 + idx)  # B, C, D, etc.
                                            chart.add_series({
                                                'name': 'Dirty Profile',
                                                'categories': f'=\'Roll Profile\'!$A${data_start_row+1}:$A${data_end_row+1}',
                                                'values': f'=\'Roll Profile\'!${col_letter}${data_start_row+1}:${col_letter}${data_end_row+1}',
                                                'line': {'color': '#1f77b4' if idx == 0 else None, 'width': 2.5},
                                                'marker': {
                                                    'type': 'circle', 
                                                    'size': 7,
                                                    'fill': {'color': '#1f77b4' if idx == 0 else None}
                                                },
                                            })
                                        
                                        chart.set_title({'name': f'{roll_id}', 'name_font': {'size': 14, 'bold': True}})
                                        chart.set_x_axis({
                                            'name': 'Distance (mm)',
                                            'name_font': {'size': 11, 'bold': True},
                                            'num_font': {'size': 10}
                                        })
                                        chart.set_y_axis({
                                            'name': 'Diameter (mm)',
                                            'name_font': {'size': 11, 'bold': True},
                                            'num_font': {'size': 10}
                                        })
                                        chart.set_size({'width': 720, 'height':350})
                                        chart.set_legend({'position': 'right', 'font': {'size': 10}})
                                        chart.set_style(10)
                                        
                                        # Insert chart
                                        worksheet.insert_chart(f'E{start_row}', chart)
                                        
                                        # Adjust column widths
                                        worksheet.set_column('A:A', 12)
                                        for i in range(len(dates_list)):
                                            worksheet.set_column(i+1, i+1, 15)
                                
                                except ImportError:
                                    # Fallback to openpyxl without chart
                                    summary_df = pd.DataFrame({
                                        'Roll No': [roll_id],
                                        'Date(s)': [', '.join(dates_selected)]
                                    })
                                    
                                    pivot_data = plot_data.pivot_table(
                                        index='Distance', 
                                        columns='DateLabel', 
                                        values='Diameter',
                                        aggfunc='first'
                                    ).reset_index()
                                    
                                    with pd.ExcelWriter(output, engine='openpyxl') as writer:
                                        summary_df.to_excel(writer, sheet_name='Roll Profile', index=False, startrow=0)
                                        pivot_data.to_excel(writer, sheet_name='Roll Profile', index=False, startrow=3)
                                        plot_data.to_excel(writer, sheet_name='Raw Data', index=False)
                                
                                output.seek(0)
                                return output.getvalue()

                            st.download_button(
                                "⬇️ Download Chart as Excel",
                                data=to_chart_excel_bytes(plot_df, selected_roll, chosen_dates),
                                file_name=f"roll_profile_{selected_roll}.xlsx",
                                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                                use_container_width=True
                            )
            else:
                st.info("Please choose a Roll No from the dropdown to plot.")

st.markdown('</div>', unsafe_allow_html=True)
<!DOCTYPE html>
<html>
<head>
    <title>Pinch Roll Data Form</title>
</head>
<body>
    <p>
        <a href="https://rollprofile.streamlit.app/">Click here for Pinch Roll Data Form</a>
    </p>
</body>
</html>











